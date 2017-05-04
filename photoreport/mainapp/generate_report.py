from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt
from docx.shared import Inches
import os
from django.conf import settings
from PIL import Image

class DocumentGenerator(object):

    def create(self, project):
        document = Document()
        obj_style = document.styles['Normal']
        obj_font = obj_style.font
        obj_font.size = Pt(12)
        obj_font.name = 'Tahoma'

        img_idx = 1
        for image in project.sorted_images:
            if len(image.caption) > 5: #no caption if less than 5 characters in caption
	    	img_txt = str(img_idx) + ". " + image.caption
            	p = document.add_paragraph(style ='Normal',text = img_txt)
		p = document.paragraphs[-1]
		p.line_spacing = 1.0
            	#p = document.add_paragraph(style ='ListBullet',text = img_txt)
		#p.num_id = img_idx 
		#p.level = 0
		#p.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
		p.line_spacing_rule = WD_LINE_SPACING.EXACTLY
	
		#p.line_spacing = 1.5
		

            	img_full_url = os.path.join(
                	settings.MEDIA_ROOT, image.image.url.split('/media/')[1]
            	)
		im = Image.open(img_full_url)
		(img_width, img_len) = im.size
            	if img_width > img_len:   #wide picture
			document.add_picture(img_full_url, height=Inches(3.25))
            	else: 
			document.add_picture(img_full_url, width=Inches(2.44))
		
		p = document.paragraphs[-1]  # grab the last paragraph
            	p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # center image
		img_idx += 1
            	if img_idx % 2:#add two images and then page break
                	document.add_page_break()
            	#else:
                #	document.add_paragraph(" ") #add line break
        return document
